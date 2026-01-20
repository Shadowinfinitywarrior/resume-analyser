import csv
import os
import uuid
import io
from datetime import datetime
import pypdf
import docx

DATA_DIR = 'data'

def get_next_id(filename):
    path = os.path.join(DATA_DIR, filename)
    if not os.path.exists(path):
        return 1
    with open(path, 'r', newline='', encoding='utf-8') as f:
        reader = csv.reader(f)
        data = list(reader)
        if len(data) <= 1:
            return 1
        # Find the last row with a valid ID
        for row in reversed(data[1:]):  # Skip header, start from bottom
            if row and row[0] and row[0].strip():  # Check if ID exists and is not empty
                try:
                    return int(row[0]) + 1
                except ValueError:
                    continue
        return 1

def load_csv(filename):
    path = os.path.join(DATA_DIR, filename)
    if not os.path.exists(path):
        return []
    with open(path, 'r', newline='', encoding='utf-8') as f:
        reader = list(csv.DictReader(f))
        return reader

def append_csv(filename, fieldnames, row_dict):
    path = os.path.join(DATA_DIR, filename)
    with open(path, 'a', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writerow(row_dict)

from werkzeug.security import generate_password_hash, check_password_hash

# --- USER MANAGEMENT ---
def load_users():
    return load_csv('users.csv')

def save_user(username, password, role, email):
    # Check if email already exists
    users = load_users()
    if any(u['email'] == email for u in users):
        return False

    new_id = get_next_id('users.csv')
    hashed_password = generate_password_hash(password)
    
    append_csv('users.csv', ['id', 'username', 'password', 'role', 'email'], {
        'id': new_id,
        'username': username,
        'password': hashed_password,
        'role': role,
        'email': email
    })
    return True

def authenticate_user(email, password):
    users = load_users()
    for user in users:
        if user['email'] == email:
            if check_password_hash(user['password'], password):
                return user
    return None

def initialize_admin():
    users = load_users()
    # Check if admin exists
    if not any(u['role'] == 'admin' for u in users):
        print("Initializing default admin...")
        save_user('Admin', 'admin123', 'admin', 'admin@resume.com')

def delete_user(user_id):
    users = load_csv('users.csv')
    new_users = [u for u in users if str(u['id']) != str(user_id) and u['role'] != 'admin'] # Prevent admin deletion
    
    if len(users) == len(new_users):
        return False
        
    path = os.path.join(DATA_DIR, 'users.csv')
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=['id', 'username', 'password', 'role', 'email'])
        writer.writeheader()
        writer.writerows(new_users)
    return True

# --- JOB MANAGEMENT ---
def load_jobs():
    return load_csv('jobs.csv')

def get_job_by_id(job_id):
    jobs = load_jobs()
    for job in jobs:
        if str(job['id']) == str(job_id):
            return job
    return None

def save_job(hr_id, title, description, skills, vacancies):
    new_id = get_next_id('jobs.csv')
    append_csv('jobs.csv', ['id', 'hr_id', 'title', 'description', 'skills_required', 'vacancies', 'status'], {
        'id': new_id,
        'hr_id': hr_id,
        'title': title,
        'description': description,
        'skills_required': skills,
        'vacancies': vacancies,
        'status': 'Open'
    })
    return True

def delete_job(job_id):
    jobs = load_jobs()
    new_jobs = [j for j in jobs if str(j['id']) != str(job_id)]
    
    if len(jobs) == len(new_jobs):
        return False
        
    path = os.path.join(DATA_DIR, 'jobs.csv')
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=['id', 'hr_id', 'title', 'description', 'skills_required', 'vacancies', 'status'])
        writer.writeheader()
        writer.writerows(new_jobs)
        
    # Also delete associated applications
    apps = load_csv('applications.csv')
    new_apps = [a for a in apps if str(a['job_id']) != str(job_id)]
    
    path_apps = os.path.join(DATA_DIR, 'applications.csv')
    with open(path_apps, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=['id', 'job_id', 'user_id', 'resume_id', 'status', 'hr_notes', 'score'])
        writer.writeheader()
        writer.writerows(new_apps)
        
    return True

def update_job_status(job_id, new_status):
    jobs = load_jobs()
    updated = False
    new_jobs = []
    
    for job in jobs:
        if str(job['id']) == str(job_id):
            job['status'] = new_status
            updated = True
        new_jobs.append(job)
        
    if updated:
        path = os.path.join(DATA_DIR, 'jobs.csv')
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=['id', 'hr_id', 'title', 'description', 'skills_required', 'vacancies', 'status'])
            writer.writeheader()
            writer.writerows(new_jobs)
            
    return updated

# --- RESUME START ---
def extract_text_from_pdf(file_stream):
    try:
        file_stream.seek(0)
        reader = pypdf.PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_stream):
    try:
        if isinstance(file_stream, (str, bytes, os.PathLike)):
            doc = docx.Document(file_stream)
        else:
            file_stream.seek(0)
            doc = docx.Document(file_stream)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def save_resume(user_id, filename, file_storage):
    new_id = get_next_id('resumes.csv')
    
    # Extract text based on extension
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    content_text = ""
    
    if ext == 'pdf':
        content_text = extract_text_from_pdf(file_storage)
    elif ext in ['docx', 'doc']:
        content_text = extract_text_from_docx(file_storage)
    else:
        # Assume text/plain or try utf-8
        try:
            file_storage.seek(0)
            content_text = file_storage.read().decode('utf-8', errors='ignore')
        except:
            content_text = "Could not parse file content."
            
    # Save the original file for proper rendering
    uploads_dir = os.path.join(DATA_DIR, 'uploads')
    os.makedirs(uploads_dir, exist_ok=True)
    
    # Use UUID to prevent filename collisions
    saved_filename = f"{uuid.uuid4()}_{filename}"
    file_path = os.path.join(uploads_dir, saved_filename)
    
    file_storage.seek(0)
    file_storage.save(file_path)
    
    # If binary read moved cursor, reset might be needed, but pypdf/docx usually handle stream.
    # However, saving the raw file might be useful in a real app, but here we only save text to CSV.
    
    append_csv('resumes.csv', ['id', 'user_id', 'filename', 'content_text', 'upload_date', 'file_path'], {
        'id': new_id,
        'user_id': user_id,
        'filename': filename,
        'content_text': content_text.replace('\r', ''), # Clean up for CSV, keep \n
        'upload_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'file_path': file_path
    })
    return new_id

def get_user_resumes(user_id):
    resumes = load_csv('resumes.csv')
    return [r for r in resumes if str(r['user_id']) == str(user_id)]

def get_user_resume(user_id):
    resumes = load_csv('resumes.csv')
    user_resumes = [r for r in resumes if str(r['user_id']) == str(user_id)]
    if user_resumes:
        return user_resumes[-1]
    return None

def get_resume_by_id(resume_id):
    resumes = load_csv('resumes.csv')
    for r in resumes:
        if str(r['id']) == str(resume_id):
            return r
    return None

# --- APPLICATION MANAGEMENT ---
def save_application(job_id, user_id, resume_id, score, eligibility='Low'):
    new_id = get_next_id('applications.csv')
    # Check if already applied
    existing = load_csv('applications.csv')
    for app in existing:
        if str(app['job_id']) == str(job_id) and str(app['user_id']) == str(user_id):
            return False 

    append_csv('applications.csv', ['id', 'job_id', 'user_id', 'resume_id', 'status', 'hr_notes', 'score', 'eligibility'], {
        'id': new_id,
        'job_id': job_id,
        'user_id': user_id,
        'resume_id': resume_id,
        'status': 'Applied',
        'hr_notes': '',
        'score': score,
        'eligibility': eligibility
    })
    return True

def get_user_applications(user_id):
    apps = load_csv('applications.csv')
    jobs = {j['id']: j['title'] for j in load_jobs()}
    user_apps = []
    for app in apps:
        if str(app['user_id']) == str(user_id):
            app['job_title'] = jobs.get(app['job_id'], 'Unknown Job')
            user_apps.append(app)
    return user_apps

# --- AI SIMULATION ---

# --- AI SIMULATION ---
def class_based_compatibility(resume_text, job_skills):
    if not job_skills: return 0, []
    resume_text = resume_text.lower()
    skills = [s.strip().lower() for s in job_skills.split(',')]
    matched = [s for s in skills if s in resume_text]
    score = (len(matched) / len(skills)) * 100 if skills else 0
    return round(score, 2), matched

import re

def deep_resume_analysis(text):
    text_lower = text.lower()
    score = 100
    suggestions = []
    
    # 1. Section Check
    required_sections = {
        'Education': ['education', 'academic', 'degree', 'university', 'college'],
        'Experience': ['experience', 'work history', 'employment', 'internship'],
        'Skills': ['skills', 'technologies', 'competencies', 'proficiencies'],
        'Projects': ['projects', 'undertakings', 'portfolio'],
        'Contact': ['email', 'phone', 'contact', 'address']
    }
    
    missing_sections = []
    for section, keywords in required_sections.items():
        if not any(k in text_lower for k in keywords):
            missing_sections.append(section)
            score -= 10
            
    if missing_sections:
        suggestions.append(f"Missing important sections: {', '.join(missing_sections)}")
    else:
        suggestions.append("Structure: All key sections detected. Good job!")

    # 2. Contact Info Check (Regex)
    # Simple email regex
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    if not re.search(email_pattern, text):
        suggestions.append("Critical: No email address found. Recruiters cannot contact you.")
        score -= 20
        
    # Simple phone regex (very basic)
    phone_pattern = r'\b\d{10}\b|\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
    if not re.search(phone_pattern, text):
        suggestions.append("Warning: No clear phone number found.")
        score -= 5

    # 3. Action Verbs Check
    # A small sample list of strong action verbs
    action_verbs = ['led', 'managed', 'developed', 'created', 'implemented', 'designed', 'analyzed', 'solved', 'achieved', 'improved']
    found_verbs = [verb for verb in action_verbs if verb in text_lower]
    
    if len(found_verbs) < 3:
        suggestions.append(f"Weak language: Try using more action verbs like {', '.join(action_verbs[:5])}.")
        score -= 10
    else:
        suggestions.append(f"Language: Good use of action verbs ({len(found_verbs)} found).")
        
    # 4. Content Length (Rough approximation)
    word_count = len(text.split())
    if word_count < 200:
        suggestions.append("Length: Resume seems too short. Elaborate on your experiences.")
        score -= 10
    elif word_count > 1500:
        suggestions.append("Length: Resume might be too long. Keep it concise.")
    else:
        suggestions.append("Length: Good word count.")

    return max(0, score), suggestions

def basic_resume_analysis(text):
    # Wrapper to maintain backward compatibility if needed, or just redirect
    return deep_resume_analysis(text)

def check_job_satisfaction(resume_text, job_description, detailed=False):
    """Enhanced job satisfaction check with skill categorization and recommendations."""
    resume_lower = resume_text.lower()
    jd_lower = job_description.lower()
    
    # Extract words from JD and resume
    jd_words = re.findall(r'\b\w+\b', jd_lower)
    resume_words = set(re.findall(r'\b\w+\b', resume_lower))
    
    # Expanded stopwords
    common_stopwords = {
        'and', 'the', 'to', 'of', 'in', 'a', 'for', 'with', 'on', 'is', 'at', 'an', 'or', 'be', 'as', 'are', 'will',
        'this', 'that', 'from', 'have', 'has', 'had', 'but', 'they', 'their', 'what', 'which', 'who', 'when',
        'where', 'how', 'all', 'each', 'other', 'some', 'such', 'into', 'than', 'them', 'these', 'those',
        'would', 'should', 'could', 'been', 'being', 'were', 'was', 'can', 'may', 'must', 'shall', 'years'
    }
    
    # Extract important keywords (filtering stopwords and short words)
    important_keywords = [w for w in jd_words if w not in common_stopwords and len(w) > 3]
    unique_keywords = set(important_keywords)
    
    # Technical skill indicators
    tech_indicators = {
        'python', 'java', 'javascript', 'c++', 'sql', 'html', 'css', 'react', 'angular', 'vue',
        'django', 'flask', 'spring', 'nodejs', 'aws', 'azure', 'docker', 'kubernetes', 'git',
        'machine', 'learning', 'tensorflow', 'pytorch', 'data', 'analysis', 'science',
        'cloud', 'devops', 'api', 'rest', 'database', 'mongodb', 'postgresql', 'mysql'
    }
    
    # Categorize keywords
    matched = [w for w in unique_keywords if w in resume_words]
    missing = [w for w in unique_keywords if w not in resume_words]
    
    # Identify technical skills vs general keywords
    missing_tech = [w for w in missing if w in tech_indicators or any(tech in w for tech in tech_indicators)]
    missing_general = [w for w in missing if w not in missing_tech]
    
    if not unique_keywords:
        return 0, {"missing_keywords": [], "message": "Job Description too short to analyze."}
    
    score = (len(matched) / len(unique_keywords)) * 100
    
    if detailed:
        # Return detailed breakdown
        return round(score), {
            "matched_count": len(matched),
            "total_keywords": len(unique_keywords),
            "missing_technical": missing_tech[:5],  # Top 5 technical skills
            "missing_general": missing_general[:5],   # Top 5 general keywords
            "all_missing": sorted(list(missing))[:10],
            "eligibility_level": "High" if score >= 70 else "Medium" if score >= 40 else "Low",
            "recommendation": _get_recommendation(score, missing_tech)
        }
    else:
        # Return simple format for compatibility
        return round(score), sorted(list(missing))[:10]

def _get_recommendation(score, missing_tech_skills):
    """Generate recommendation based on score and missing skills."""
    if score >= 70:
        return "You are highly eligible for this position. Apply with confidence!"
    elif score >= 50:
        if missing_tech_skills:
            return f"You have a good foundation. Consider adding skills like: {', '.join(missing_tech_skills[:3])}"
        return "You meet most requirements. Highlight your relevant experience in your application."
    elif score >= 30:
        if missing_tech_skills:
            return f"You may need to strengthen skills in: {', '.join(missing_tech_skills[:3])}. Consider taking online courses."
        return "You meet some requirements. Emphasize transferable skills and willingness to learn."
    else:
        return "This position requires skills you may not have yet. Consider gaining experience in the key areas mentioned."

# --- ADMIN UTILS ---
def get_system_stats():
    users = load_users()
    jobs = load_jobs()
    apps = load_csv('applications.csv')
    return {
        'users': len(users),
        'jobs': len(jobs),
        'apps': len(apps),
        'user_list': users
    }
    
def get_all_jobs():
    return load_jobs()

# --- HR UTILS ---
def get_hr_jobs_with_applications(hr_id):
    jobs = load_jobs()
    my_jobs = {j['id']: j for j in jobs if str(j['hr_id']) == str(hr_id)}
    
    if not my_jobs:
        return {}
        
    apps = load_csv('applications.csv')
    users = {u['id']: u['username'] for u in load_users()}
    
    result = {}
    for jid, job in my_jobs.items():
        result[jid] = {
            'title': job['title'],
            'vacancies': job['vacancies'],
            'status': job['status'],
            'apps': []
        }
    
    for app in apps:
        if app['job_id'] in result:
            app['username'] = users.get(app['user_id'], 'Unknown')
            result[app['job_id']]['apps'].append(app)
            
    # Sort applications by score
    for jid in result:
        result[jid]['apps'].sort(key=lambda x: float(x.get('score', 0) or 0), reverse=True)
        
    return result

def update_application_status(app_id, status, notes):
    apps = load_csv('applications.csv')
    updated = False
    new_apps = []
    
    fieldnames = ['id', 'job_id', 'user_id', 'resume_id', 'status', 'hr_notes', 'score', 'eligibility']
    
    for app in apps:
        if str(app['id']) == str(app_id):
            app['status'] = status
            app['hr_notes'] = notes
            updated = True
        new_apps.append(app)
        
    if updated:
        path = os.path.join(DATA_DIR, 'applications.csv')
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(new_apps)
    return updated

def bulk_update_applications(app_ids, status, justification):
    apps = load_csv('applications.csv')
    updated_count = 0
    new_apps = []
    fieldnames = ['id', 'job_id', 'user_id', 'resume_id', 'status', 'hr_notes', 'score', 'eligibility']
    
    # Ensure app_ids is a set of strings for easy lookup
    target_ids = set(str(aid) for aid in app_ids)
    
    for app in apps:
        if str(app['id']) in target_ids:
            app['status'] = status
            # Append new justification to existing notes or set it
            current_notes = app.get('hr_notes', '')
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            app['hr_notes'] = f"{current_notes} | [{timestamp}] {status}: {justification}".strip(" |")
            updated_count += 1
        new_apps.append(app)
        
    if updated_count > 0:
        path = os.path.join(DATA_DIR, 'applications.csv')
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(new_apps)
            
    return updated_count

# --- BULK RESUME SCREENING ---
import zipfile
import tempfile
import shutil

def extract_and_parse_resumes(zip_file):
    """Extract resumes from ZIP and parse each one."""
    resumes_data = []
    
    # Create temporary directory
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extract ZIP
        with zipfile.ZipFile(zip_file, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Parse each file
        for root, dirs, files in os.walk(temp_dir):
            for filename in files:
                if filename.startswith('.'):
                    continue  # Skip hidden files
                    
                filepath = os.path.join(root, filename)
                ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
                
                content_text = ""
                
                try:
                    if ext == 'pdf':
                        with open(filepath, 'rb') as f:
                            content_text = extract_text_from_pdf(f)
                    elif ext in ['docx', 'doc']:
                        content_text = extract_text_from_docx(filepath)
                    elif ext == 'txt':
                        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                            content_text = f.read()
                    else:
                        continue  # Skip unsupported formats
                    
                    if content_text.strip():
                        resumes_data.append({
                            'filename': filename,
                            'content': content_text,
                            'original_path': filepath
                        })
                except Exception as e:
                    print(f"Error parsing {filename}: {e}")
                    continue
    
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    return resumes_data

def screen_candidates(job_id, resumes_data):
    """Screen candidates and generate recommendations."""
    job = get_job_by_id(job_id)
    if not job:
        return []
    
    job_description = f"{job['title']} {job['description']} {job['skills_required']}"
    vacancies = int(job.get('vacancies', 1))
    
    # Score each resume
    scored_candidates = []
    for resume in resumes_data:
        score, details = check_job_satisfaction(resume['content'], job_description, detailed=True)
        
        # Determine recommendation
        if score >= 70:
            recommendation = "Select"
            justification = f"Strong match ({score}%). Candidate possesses most required skills."
        elif score >= 50:
            recommendation = "Review"
            justification = f"Moderate match ({score}%). Consider for interview to assess fit."
        else:
            recommendation = "Reject"
            missing_skills = ', '.join(details.get('missing_technical', [])[:3])
            justification = f"Low match ({score}%). Missing key skills: {missing_skills or 'multiple areas'}."
        
        scored_candidates.append({
            'filename': resume['filename'],
            'content': resume['content'],
            'score': score,
            'recommendation': recommendation,
            'justification': justification,
            'details': details,
            'original_path': resume.get('original_path')
        })
    
    # Sort by score descending
    scored_candidates.sort(key=lambda x: x['score'], reverse=True)
    
    # Auto-select top candidates based on vacancies
    for i, candidate in enumerate(scored_candidates):
        if i < vacancies and candidate['score'] >= 60:
            if candidate['recommendation'] != 'Select':
                candidate['recommendation'] = 'Select'
                candidate['justification'] = f"Top {vacancies} candidate ({candidate['score']}%). Meets requirements."
    
    return scored_candidates

def save_candidate_to_pool(job_id, candidate_data):
    """Save screened candidate to candidate pool."""
    new_id = get_next_id('candidate_pool.csv')
    
    # Save file to uploads
    uploads_dir = os.path.join(DATA_DIR, 'uploads')
    os.makedirs(uploads_dir, exist_ok=True)
    
    saved_filename = f"{uuid.uuid4()}_{candidate_data['filename']}"
    file_path = os.path.join(uploads_dir, saved_filename)
    
    # Copy from temp to permanent
    if candidate_data.get('original_path') and os.path.exists(candidate_data['original_path']):
        shutil.copy(candidate_data['original_path'], file_path)
    
    append_csv('candidate_pool.csv', 
               ['id', 'job_id', 'filename', 'content_text', 'score', 'recommendation', 'justification', 'hr_decision', 'upload_date', 'file_path'],
               {
                   'id': new_id,
                   'job_id': job_id,
                   'filename': candidate_data['filename'],
                   'content_text': candidate_data['content'].replace('\r', ''),
                   'score': candidate_data['score'],
                   'recommendation': candidate_data['recommendation'],
                   'justification': candidate_data['justification'],
                   'hr_decision': '',
                   'upload_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                   'file_path': file_path
               })
    return new_id

def get_candidate_pool(job_id):
    """Get all candidates from pool for a specific job."""
    candidates = load_csv('candidate_pool.csv')
    return [c for c in candidates if str(c['job_id']) == str(job_id)]

def update_candidate_decision(candidate_id, decision, notes=''):
    """Update HR decision for a candidate."""
    candidates = load_csv('candidate_pool.csv')
    updated = False
    new_candidates = []
    
    fieldnames = ['id', 'job_id', 'filename', 'content_text', 'score', 'recommendation', 'justification', 'hr_decision', 'upload_date', 'file_path']
    
    for candidate in candidates:
        if str(candidate['id']) == str(candidate_id):
            candidate['hr_decision'] = decision
            if notes:
                candidate['justification'] = f"{candidate['justification']} | HR: {notes}"
            updated = True
        new_candidates.append(candidate)
    
    if updated:
        path = os.path.join(DATA_DIR, 'candidate_pool.csv')
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(new_candidates)
    
    return updated

# --- ADMIN MONITORING ---
def get_recent_activity():
    """Get recent activity for admin dashboard."""
    apps = load_csv('applications.csv')
    resumes = load_csv('resumes.csv')
    jobs = load_jobs()
    users = {u['id']: u['username'] for u in load_users()}
    job_titles = {j['id']: j['title'] for j in jobs}
    
    # Get latest applications
    recent_apps = sorted(apps, key=lambda x: x.get('id', '0'), reverse=True)[:10]
    for app in recent_apps:
        app['username'] = users.get(app.get('user_id'), 'Unknown')
        app['job_title'] = job_titles.get(app.get('job_id'), 'Unknown Job')
    
    # Get latest resumes
    recent_resumes = sorted(resumes, key=lambda x: x.get('upload_date', ''), reverse=True)[:10]
    for resume in recent_resumes:
        resume['username'] = users.get(resume.get('user_id'), 'Unknown')
    
    # Get latest jobs
    recent_jobs = sorted(jobs, key=lambda x: x.get('id', '0'), reverse=True)[:5]
    
    return {
        'applications': recent_apps,
        'resumes': recent_resumes,
        'jobs': recent_jobs
    }

def get_system_metrics():
    """Calculate system health metrics."""
    apps = load_csv('applications.csv')
    jobs = load_jobs()
    
    # Calculate average application score
    scores = [float(app.get('score', 0)) for app in apps if app.get('score')]
    avg_score = sum(scores) / len(scores) if scores else 0
    
    # Find most active jobs (by application count)
    job_app_counts = {}
    for app in apps:
        job_id = app.get('job_id')
        job_app_counts[job_id] = job_app_counts.get(job_id, 0) + 1
    
    most_active = sorted(job_app_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    job_titles = {j['id']: j['title'] for j in jobs}
    
    active_jobs = [{
        'title': job_titles.get(job_id, 'Unknown'),
        'count': count
    } for job_id, count in most_active]
    
    return {
        'avg_application_score': round(avg_score, 1),
        'total_applications': len(apps),
        'active_jobs': active_jobs,
        'open_positions': len([j for j in jobs if j.get('status') == 'Open'])
    }

